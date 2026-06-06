"""
Resume Text Extractor — Phase 2

Extracts and normalizes raw text from PDF and DOCX files.
Used as the first step in the AI screening pipeline.
"""

import io
import re
import unicodedata
from typing import Literal


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber (primary) with pymupdf fallback."""
    text = ""

    # Primary: pdfplumber (better for structured/tabular PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return _normalize(text)
    except Exception:
        pass

    # Fallback: pymupdf (fitz) — handles scanned/complex PDFs better
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        if text.strip():
            return _normalize(text)
    except Exception:
        pass

    raise ValueError("Unable to extract text from PDF. The file may be scanned or encrypted.")


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return _normalize("\n".join(paragraphs))
    except Exception as exc:
        raise ValueError(f"Unable to extract text from DOCX: {exc}") from exc


# ---------------------------------------------------------------------------
# Text normalization
# ---------------------------------------------------------------------------

def _normalize(text: str) -> str:
    """
    Normalize extracted text:
    - Decode unicode characters
    - Collapse excessive whitespace
    - Remove non-printable control characters
    - Preserve meaningful newlines
    """
    # Normalize unicode (e.g., ligatures, smart quotes)
    text = unicodedata.normalize("NFKD", text)

    # Remove non-printable chars except newline and tab
    text = re.sub(r"[^\S\n\t]+", " ", text)           # collapse spaces
    text = re.sub(r"\n{3,}", "\n\n", text)             # max 2 consecutive newlines
    text = re.sub(r"[ \t]+", " ", text)               # collapse inline whitespace
    text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f]", "", text)  # strip control chars

    return text.strip()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

FileType = Literal["pdf", "docx"]


def extract_text(
    file_bytes: bytes,
    file_type: FileType,
) -> str:
    """
    Extract and normalize text from a resume file.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        file_type:  "pdf" or "docx".

    Returns:
        Normalized plain text string.

    Raises:
        ValueError: If extraction fails or file type is unsupported.
    """
    ft = file_type.lower().strip(".").strip()

    if ft == "pdf":
        return _extract_pdf(file_bytes)
    elif ft in ("docx", "doc"):
        return _extract_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: '{file_type}'. Only PDF and DOCX are supported.")


def detect_file_type(filename: str, content_type: str = "") -> FileType:
    """
    Detect file type from filename extension or MIME type.
    Returns 'pdf' or 'docx'. Raises ValueError if unsupported.
    """
    name_lower = filename.lower()
    if name_lower.endswith(".pdf") or "pdf" in content_type:
        return "pdf"
    if name_lower.endswith(".docx") or name_lower.endswith(".doc") or "word" in content_type:
        return "docx"
    raise ValueError(
        f"Unsupported file type for '{filename}'. Only PDF and DOCX are supported."
    )
